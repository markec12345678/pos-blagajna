"""
Generira analizo repozitorija NutrixPOS kot .docx dokument.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = "/home/z/my-project/download/analiza_nutrixpos.docx"

doc = Document()

# --- Osnovne nastavitve pisave ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.3
style.paragraph_format.space_after = Pt(6)

# --- Robovi strani ---
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(it)


def add_table_row(table, cells):
    row = table.add_row()
    for i, txt in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.add_run(str(txt))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ============================================================================
# NASLOVNICA
# ============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ANALIZA REPOZITORIJA")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("NutrixPOS — sistem za upravljanje prodaje (POS)")
r.font.size = Pt(16)
r.italic = True
r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("GitHub: https://github.com/nutrixpos/pos.git\n")
r.font.size = Pt(11)
r = meta.add_run("Datum analize: 21. julij 2026\n")
r.font.size = Pt(11)
r = meta.add_run("Jezik dokumenta: Slovenščina")
r.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

warn = doc.add_paragraph()
warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = warn.add_run("⚠ VARNOSTNO OPOZORILO\n")
r.bold = True
r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
r.font.size = Pt(14)
r = warn.add_run(
    "V prvotnem sporočilu je bil izpostavljen GitHub Personal Access Token. "
    "Žeton je treba takoj razveljaviti na https://github.com/settings/tokens."
)
r.font.size = Pt(10)
r.italic = True

doc.add_page_break()

# ============================================================================
# KAZALO VSEBINE (ročno)
# ============================================================================
add_heading("Kazalo vsebine", level=1)
toc_items = [
    "1. Povzetek projekta",
    "2. Tehnološki sklad",
    "3. Arhitektura in struktura",
    "4. Backend analiza (Go)",
    "5. Frontend analiza (Vue 3)",
    "6. Varnostna analiza",
    "7. Kakovost kode in vzdrževalnost",
    "8. Infrastruktura in CI/CD",
    "9. Močne strani",
    "10. Šibke točke in tveganja",
    "11. Priporočila za izboljšave",
    "12. Sklepna ocena",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.5)

doc.add_page_break()

# ============================================================================
# 1. POVZETEK
# ============================================================================
add_heading("1. Povzetek projekta", level=1)

add_para(
    "NutrixPOS je odprtokodni sistem za upravljanje prodaje (Point of Sale) "
    "namenjen predvsem restavracijam in maloprodajnim trgovinam. Omogoča "
    "upravljanje zalog, prodaje, izdelkov, naročil, kupcev in računov. "
    "Projekt je aktivno v razvoju in, kot opozarja README, ne zagotavlja "
    "zadržne združljivosti do izdaje stabilne različice."
)

add_para("Ključni atributi projekta:", bold=True)
add_bullets([
    "Ime: NutrixPOS (modul se v klicih imenuje 'nutrix')",
    "Licenca: GPL (18 KB datoteka LICENSE)",
    "Status: Aktivni razvoj, v0.5.5 (po docker-compose)",
    "Število avtorjev: 1 (elmawardy) — samostojni projekt",
    "Število Go datotek: 66 (~11.500 vrstic)",
    "Število frontend datotek: 47 (~10.100 vrstic)",
    "Testi: 0 (nobenih _test.go ali .spec.ts datotek)",
])

# ============================================================================
# 2. TEHNOLOŠKI SKLAD
# ============================================================================
add_heading("2. Tehnološki sklad", level=1)

add_para(
    "Projekt je monorepo z ločenima backend (Go) in frontend (Vue 3) "
    "sklopoma. Uporablja MongoDB kot primarno podatkovno bazo, z dodiplomskim "
    "podprtjem za Zitadel (OIDC) kot alternativni avtentikacijski ponudnik."
)

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Sklad"
hdr[1].text = "Tehnologija"
hdr[2].text = "Različica / opomba"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
    shade_cell(c, "1F3A5F")
    for p in c.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows = [
    ("Backend jezik", "Go", "1.25.0 (go.mod) / 1.24.4 (Dockerfile) — neskladje"),
    ("Spletni okvir", "gorilla/mux", "v1.8.1"),
    ("Podatkovna baza", "MongoDB", "mongo-driver v1.16.1, MongoDB 5.0 v docker-compose"),
    ("CLI / konfig", "cobra + viper", "v1.8.1 / v1.19.0"),
    ("Loggerji", "zerolog + zap", "Dva loggerja — neskladje"),
    ("Avtentikacija", "JWT (golang-jwt v5) + Zitadel OIDC", "Pogojno vključena"),
    ("WebSocket", "olahol/melody", "Realnočasovna obvestila"),
    ("Tiskalniki", "elmawardy/escpos + chromedp", "Toplotni tiskalnik + PDF tisk"),
    ("Frontend okvir", "Vue 3", "v3.5.13 + TypeScript 5.8"),
    ("UI knjižnica", "PrimeVue + PrimeFlex", "v4.3.x"),
    ("State management", "Pinia", "v3.0.1"),
    ("HTTP odjemalec", "fetch + axios", "Mešanje dveh pristopov"),
    ("Grafi", "Chart.js + vue-chartjs", "v4.4.9 / v5.3.2"),
    ("i18n", "vue-i18n", "ar, en, pt-br"),
    ("Gradnja", "Vite 6", "z vite-plugin-singlefile"),
    ("E2E testiranje", "Cypress 14", "konfigurirano, vendar brez testov"),
    ("E2E enote", "Vitest 3", "konfigurirano, vendar brez testov"),
    ("Pakiranje", "NSIS (.exe) za Windows", "Preko GitHub Actions"),
]
for r0 in rows:
    add_table_row(table, r0)

doc.add_paragraph()

# ============================================================================
# 3. ARHITEKTURA
# ============================================================================
add_heading("3. Arhitektura in struktura", level=1)

add_para(
    "Backend sledi modularni monolitni arhitekturi z vtičniškim sistemom "
    "za module. Vsak modul implementira vmesnik IBaseModule in lahko "
    "neobvezno implementira IHttpModule, IBackgroundWorkerModule ter "
    "ISeederModule. AppManager skrbi za registracijo, zagon in zaključek "
    "modulov."
)

add_para("Trenutno sta registrirana dva modula:", bold=True)
add_bullets([
    "core — glavna poslovna logika (materiali, izdelki, naročila, prodaja, kupci, kategorije, nastavitve, jeziki, obvestila)",
    "hubsync — sinhronizacija z oddaljenim 'hub' API-jem (uporablja bearer token in intervalno sinhronizacijo, privzeto 60 s)",
])

add_para("Glavne mape in njihov pomen:", bold=True)
add_bullets([
    "/cmd — CLI vhodi (root.go z Execute() in seed.go)",
    "/common — skupne rešitve: config (viper), logger (zerolog + zap), database (singleton), userio, helpers, customerrors",
    "/modules/core — modeli, servisi, handlerji, middlewares (CORS), DTO-ji, API specifikacija",
    "/modules/auth — middlewares (JWT, bcrypt, Zitadel), modeli uporabnikov, handlerji za prijavo/registracijo",
    "/modules/hubsync — sinhronizacija z centralnim API-jem",
    "/modules/imodules.go, appmanager.go, modulebuilder.go, backgroundworkers.go — infrastruktura modulov",
    "/frontend — Vue 3 aplikacija (pages, components, stores, services)",
    "/assets — predloge računov (Handlebars) in jezikovni paketi (JSON)",
])

add_para(
    "Aplikacijski vhod (main.go) inicializira logger, konfiguracijo in "
    "prompter (BubbleTea TUI), nato prenese nadzor RootProcess.Execute(). "
    "Če konfiguracija nima določenega podatkovnega gostitelja, se zažene "
    "setup strežnik na :8000, ki omogoča konfiguracijo preko brskalnika "
    "(Setup.vue), nato pa se proces znova zažene z novo konfiguracijo.",
    italic=False,
)

# ============================================================================
# 4. BACKEND ANALIZA
# ============================================================================
add_heading("4. Backend analiza (Go)", level=1)

add_heading("4.1 Zagon in konfiguracija", level=2)
add_para(
    "Konfiguracija se bere iz config.yaml z Viperjem, podprta pa so tudi "
    "okoljska spremenljivka (npr. DATABASES_0_HOST). Privzeta konfiguracija "
    "vključuje podatke o bazi, JWT skrivnost, pot do Zitadel ključa, pot "
    "do mape za uploads, časovni pas in okolje (dev/prod). Setup strežnik "
    "lahko zapiše config.yaml na podlagi POST zahtevka — to je kritična "
    "točka, ki jo obravnavamo v varnostnem poglavju."
)

add_heading("4.2 Registracija HTTP poti", level=2)
add_para(
    "Vse HTTP poti so ročno registrirane v core.RegisterHttpHandlers() "
    "z veriženjem middlewarov AllowCors → AllowAnyOfRoles → Handler. "
    "Registriranih je približno 50+ endpointov, ki pokrivajo: "
    "avtentikacijo, kupce, material, vnose materiala, kategorije, "
    "naročila (z vsemi prehodmi stanj), izdelke, nastavitve, jezike, "
    "odpise (disposals) in WebSocket (/ws)."
)

add_para("Vloge, uporabljene v sistemu:", bold=True)
add_bullets([
    "superuser — popoln dostop (prvi uporabnik ob registraciji)",
    "admin — upravljanje inventarja, izdelkov, prodaje, kupcev",
    "cashier — prodaja, naročila, kupci",
    "chef — kuhinja, materiali, naročila",
])

add_heading("4.3 Vzorec handlerjev", level=2)
add_para(
    "Handlerji so implementirani kot tovarniške funkcije, ki sprejmejo "
    "config in logger ter vrnejo http.HandlerFunc (npr. "
    "handlers.GetMaterials(c.Config, c.Logger)). Ta pristop pomeni, da se "
    "ob vsaki registraciji route kreira nova zaprtja (closure), kar je "
    "pomnilniško potratno pri ~50 registracijah, vendar ne kritično."
)

add_heading("4.4 Povezava s podatkovno bazo", level=2)
add_para(
    "Uporablja se singleton vzorec v common/database.go z double-checked "
    "locking. To je dobra praksa, saj preprečuje odpiranje več povezav. "
    "AGENTS.md eksplicitno opozarja, da razvijalci ne smejo uporabljati "
    "mongo.Connect() direktno. V dev načinu je timeout 1000 sekund "
    "(nenavadno dolgo), v produkciji 5 sekund."
)

add_heading("4.5 Modeli in poslovna logika", level=2)
add_para(
    "Glavne entitete so: Material (sinonim: Component, Inventory Item), "
    "Product (sinonim: Recipe), Order, OrderItem, MaterialEntry, "
    "SalesLogs, Category, Customer, Disposal. Modeli imajo tri tage "
    "(json, bson, mapstructure), kar omogoča istočasno JSON serializacijo, "
    "MongoDB shranjevanje in Viper mapiranje."
)

add_para(
    "Posebnost: JSONFloat tip (float64) marshal-a neskončno v -1, kar je "
    "kvirk za obravnavo deljenja z nič pri izračunu stroškov. Material "
    "ima tudi koncept MaterialEntry (posamezen nakup) z lastnostmi, kot "
    "so SKU, rok trajanja in dobavitelj. Cena se izračuna na več načinov: "
    "povprečna cena (avgcost), točna cena (cost po vnosu) in fiksni strošek."
)

add_heading("4.6 Sinhronizacija (hubsync)", level=2)
add_para(
    "HubSync modul v ozadju vsakih 60 sekund (konfigurabilno) sinhronizira "
    "material, izdelke, prodajo in odpise z oddaljenim 'hub' API-jem z "
    "uporabo bearer tokena. To omogoča povezovanje več POS lokacij s "
    "centralnim sistemom — korak h proti franšiznim ali večlokacijskim "
    "scenarijem."
)

# ============================================================================
# 5. FRONTEND ANALIZA
# ============================================================================
add_heading("5. Frontend analiza (Vue 3)", level=1)

add_para(
    "Frontend je klasična Vue 3 SPA aplikacija z TypeScriptom, zgrajena "
    "z Vite 6. UI temelji na PrimeVue komponentah s PrimeFlex utilitami. "
    "State se upravlja s Pinia, avtentikacijsko stanje pa je shranjeno v "
    "localStorage (kar je varnostno tvegano — glej poglavje 6)."
)

add_heading("5.1 Struktura", level=2)
add_bullets([
    "/pages — 16 Vue komponent za posamezne strani (Home, Login, Inventory, Sales, Products, Orders, Customers, Kitchen, Settings, Users, Hubsync, Languages, Setup, AdminSetup, Profile, NoAccessView)",
    "/components — 18 ponovno uporabnih komponent (OrderView, MealCard, PickProduct, PickMaterial, PickCustomer, ...)",
    "/classes — TypeScript razredi (Order, OrderItem, Notification)",
    "/services — api.ts (fetch wrapper), auth.ts (avtentikacija), upload.ts",
    "/stores — Pinia store",
    "/router — Vue Router s preverjanjem vlog na vsaki poti",
])

add_heading("5.2 Avtentikacija in avtorizacija", level=2)
add_para(
    "JWT token se po prijavi shrani v localStorage pod ključem "
    "'nutrix_token' in uporabniški objekt pod 'nutrix_user'. reactive() "
    "ref-ova (accessToken, currentUser, isAuthenticated) omogočajo "
    "reaktivno posodabljanje UI-ja. Router na vsaki poti preverja vloga "
    "in preusmerja neavtorizirane uporabnike na /login ali /no-access."
)

add_para(
    "Pomembna napaka: funkcija register() uporablja relativno pot "
    "'/api/auth/register' brez VITE_APP_BACKEND_HOST in prefixa, medtem ko "
    "login() uporablja popolno URL s spremenljivkami. To pomeni, da "
    "registracija verjetno ne deluje, ko je frontend ločen od backend-a "
    "(npr. v Docker set-upu).",
    italic=True,
)

add_heading("5.3 Velikost in kompleksnost strani", level=2)
add_para(
    "Najobsežnejše Vue komponente: Home.vue (1417 vrstic), Products.vue "
    "(586), Inventory.vue (521), Sales.vue (445), Setup.vue (421). "
    "Komponenta Home.vue je kritično velika — tipičen 'code smell', ki "
    "zahteva razčlenitev na manjše podkomponente."
)

# ============================================================================
# 6. VARNOSTNA ANALIZA
# ============================================================================
add_heading("6. Varnostna analiza", level=1)

add_para(
    "Varnost je najšibkejši del projekta. Najdene so bile številne "
    "težave, od katerih so nekatere kritične za produkcijsko uporabo."
)

add_heading("6.1 Kritične težave", level=2)

issues_table = doc.add_table(rows=1, cols=3)
issues_table.style = "Light Grid Accent 1"
hdr = issues_table.rows[0].cells
hdr[0].text = "Težava"
hdr[1].text = "Tveganje"
hdr[2].text = "Opis"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(c, "8B0000")

issues = [
    ("Trdo kodiran JWT skrivnost",
     "Kritično",
     "config.example.yaml vsebuje 'your-super-secret-jwt-key-change-in-production'. Če uporabnik ne spremeni, lahko napadalec ponaredi JWT."),
    ("Setup endpoint zapiše config.yaml",
     "Kritično",
     "/api/setup/config POST zapiše datoteko na disk brez avtentikacije. Če je port 8000 izpostavljen, lahko napadalec prepiše konfiguracijo in preusmeri aplikacijo na svojo bazo."),
    ("CORS dovoljen povsod",
     "Visoko",
     "AllowCors middleware dovoljuje vse izvore. Primerno za razvoj, nevarno za produkcijo."),
    ("JWT v localStorage",
     "Visoko",
     "Frontend hrani JWT v localStorage. XSS napad lahko ukrade token. Priporočeno httpOnly cookie."),
    ("Brez omejevanja hitrosti",
     "Visoko",
     "Login in register endpointa nimata rate limiting. Možne brute-force napade na gesla."),
    ("Prazno geslo za MongoDB",
     "Visoko",
     "docker-compose.yaml ima prazno MONGO_INITDB_ROOT_USERNAME/PASSWORD. V produkciji je baza popolnoma odprta."),
    ("Brez HTTPS",
     "Srednje",
     "Aplikacija komunikira preko HTTP. Gesla in JWT se prenašajo v clear-text."),
    ("config.yaml v .gitignore, vendar example ima skrivnost",
     "Srednje",
     "Vzorčna konfiguracija vsebuje slab skrivnost, ki ga uporabniki pogosto pozabijo spremeniti."),
    ("Možnost injekcije preko ObjectIDFromHex",
     "Nizko",
     "Večina handlerjev pravilno uporablja primitive.ObjectIDFromHex, vendar nekateri sprejemajo niz direktno v BSON poizvedbi."),
]
for it in issues:
    add_table_row(issues_table, it)

doc.add_paragraph()

add_heading("6.2 Dobre prakse", level=2)
add_bullets([
    "Bcrypt za razprševanje gesel (pravilno uporabljen)",
    "JWT validacija preverja podpisno metodo (preprečuje alg=none napad)",
    "Vloge se preverjajo na vsakem endpointu z AllowAnyOfRoles",
    "Prvi uporabnik pri registraciji dobi superuser — pameten bootstrap",
    "Superuser-a ni mogoče izbrisati (pravilno zaščiteno)",
    "CORS preflight (OPTIONS) pravilno obravnavan na vseh endpointih",
])

# ============================================================================
# 7. KAKOVOST KODE
# ============================================================================
add_heading("7. Kakovost kode in vzdrževalnost", level=1)

add_heading("7.1 Testi", level=2)
add_para(
    "V repozitoriju ni najdena nobena testna datoteka — noben _test.go, "
    "noben .spec.ts, noben Cypress e2e test (čeprav so vsi okviri "
    "konfigurirani). AGENTS.md eksplicitno priznava: 'No tests in this "
    "repo'. To je največja šibkost projekta za vsako resno produkcijsko "
    "uporabo."
)

add_heading("7.2 Doslednost in vzorci", level=2)
add_bullets([
    "Dva loggerja (zerolog + zap) — AGENTS.md priporoča zerolog, vendar zap še vedno obstaja",
    "Go verzija: go.mod pravi 1.25.0, Dockerfile uporablja 1.24.4-alpine — neskladje",
    "Frontend meša fetch in axios za HTTP klice",
    "Registracijski URL v frontend-u je nepravilen (brez backend host)",
    "Imena funkcij vsebujejo tipkarske napake (npr. 'InesrtNewProduct' namesto 'InsertNewProduct')",
    "Komentarji v kodi so mešanica angleščine in interne dokumentacije",
    "Nekatere funkcije so prevelike (Home.vue 1417 vrstic, order.go service 1037 vrstic)",
])

add_heading("7.3 Dokumentacija", level=2)
add_para(
    "README je minimalen (810 bajtov) — samo opis, slika in opozorilo. "
    "AGENTS.md je dobro napisan tehnični vodič za AI agente z jasnim "
    "opisom arhitekture in pogostimi napakami. API specifikacija "
    "(specs.api.yaml) obstaja, vendar ni bila podrobno preverjena. "
    "Zunanja dokumentacija je na voljo na nutrixpos.com."
)

# ============================================================================
# 8. INFRASTRUKTURA
# ============================================================================
add_heading("8. Infrastruktura in CI/CD", level=1)

add_heading("8.1 Docker", level=2)
add_para(
    "docker-compose.yaml definira tri storitve: posui (frontend nginx), "
    "pos (backend) in nutrix-db (MongoDB 5.0). Zitadel (OIDC ponudnik) "
    "je zakomentiran — pripravljen za optionalno vključitev. Volumi "
    "sestavljajo ./pos_mnt za backend in ./data/mongo za podatke. "
    "Dockerfile uporablja večstopenjsko gradnjo (alpine), kar je dobra "
    "praksa."
)

add_heading("8.2 CI/CD", level=2)
add_para(
    "GitHub Actions (.github/workflows/cicd.yaml) se sproži ob objavi "
    "taga v*.*.*. Potek dela:"
)
add_bullets([
    "Gradnja frontend-a (npm install, vite build)",
    "Gradnja Go aplikacije za Windows (GOOS=windows)",
    "Prenos MongoDB MSI in ikone iz zunanjih virov",
    "Pakiranje z NSIS v Windows .exe namestilnik",
    "Objava GitHub Release z .exe priponko",
])

add_para(
    "CI/CD je omejen na Windows pakiranje. Ni avtomatiziranih testov, "
    "linting-a, signing-a ali Linux/Mac buildov. Travis/GitHub Actions "
    "ne preverita build-a na PR-ih — samo na tagih.",
    italic=True,
)

# ============================================================================
# 9. MOČNE STRANI
# ============================================================================
add_heading("9. Močne strani", level=1)

add_bullets([
    "Modularna arhitektura z jasnimi vmesniki (IBaseModule, IHttpModule, ...) — dobra razširljivost",
    "Tehnološko sodoben sklad: Go 1.25, Vue 3, TypeScript, Vite 6, Pinia",
    "Podpora za več jezikov (arabščina, angleščina, portugalščina-BR)",
    "Več avtentikacijskih strategij (internal JWT ali Zitadel OIDC)",
    "Realnočasovna obvestila preko WebSocket (melody)",
    "Tiskanje računov podprto z dvema metodama (ESC/POS toplotni tiskalnik + PDF preko chromedp)",
    "HubSync modul omogoča večlokacijsko uporabo (franchise model)",
    "Odprtokodno (GPL) — skupnostno uporabno",
    "Cross-platform backend (Go), Windows pakiranje za namizne instalacije",
    "Singleton vzorec za DB povezavo (preprečuje puščanje povezav)",
    "E2E testiranje in unit testiranje okvirja sta pripravljena (Cypress, Vitest) — samo še manjkajo testi",
])

# ============================================================================
# 10. ŠIBKE TOČKE
# ============================================================================
add_heading("10. Šibke točke in tveganja", level=1)

add_bullets([
    "Ni testov (0 datotek) — kakršnakoli sprememba je tvegana",
    "Samostojni projekt (1 avtor) — bus factor = 1",
    "Več kritičnih varnostnih težav (glej poglavje 6)",
    "Ni API validacije vhodov v večini handlerjev (samo JSON decode)",
    "Setup strežnik omogoča pisanje po disku brez avtentikacije",
    "Velike komponente (Home.vue 1417 vrstic) težko vzdržujemo",
    "Ni opolnomočene CI/CD kontrole (build samo na tagih)",
    "Ni nadzora kakovosti (linting, type-check) v CI",
    "Neskladje Go verzij (go.mod 1.25 vs Dockerfile 1.24.4)",
    "Ni HTTPS podpore vgradnji",
    "MongoDB v default docker-compose brez avtentikacije",
    "AGENTS.md priznava, da se loggerji ne uporabljajo dosledno",
    "Ni opozoril o zamenjavi JWT skrivnosti ob namestitvi",
])

# ============================================================================
# 11. PRIPOROČILA
# ============================================================================
add_heading("11. Priporočila za izboljšave", level=1)

add_heading("11.1 Takojšnje (kritično)", level=2)
add_bullets([
    "Razveljavi objavljeni GitHub Personal Access Token",
    "Odstrani trdo kodirano JWT skrivnost iz config.example.yaml in zahtevaj GENERATE_SECRET ob namestitvi",
    "Dodaj avtentikacijo ali omejitev IP-jev na /api/setup/* endpointe",
    "Konfiguriraj CORS dovoljena samo za znane izvore v produkciji",
    "Dodaj rate limiting na login/register endpointih (npr. 5 poskusov / minuto / IP)",
    "Nastavi MongoDB avtentikacijo v docker-compose in generiraj močno geslo",
])

add_heading("11.2 Kratkoročno (1–3 mesece)", level=2)
add_bullets([
    "Dodaj vsaj osnovne enote testne za backend (servisi, handlerji)",
    "Dodaj Cypress E2E test za kritične poti: login, ustvari naročilo, plačaj",
    "Počasno preidi JWT iz localStorage v httpOnly cookie",
    "Razčleni Home.vue na manjše podkomponente (cilj <300 vrstic na komponento)",
    "Popravi typo 'InesrtNewProduct' → 'InsertNewProduct'",
    "Popravi registracijski URL v frontend/services/auth.ts",
    "Poenoti HTTP odjemalca (ali fetch ali axios, ne oboje)",
    "Poenoti logger (izberi zerolog ali zap, ne oba)",
    "Posodobi Dockerfile na Go 1.25 ali go.mod na 1.24",
])

add_heading("11.3 Srednjeročno (3–6 mesecev)", level=2)
add_bullets([
    "Uvedi linting v CI (golangci-lint za Go, eslint za frontend)",
    "Dodaj build preverjanje na PR-ih (ne samo na tagih)",
    "Implementiraj HTTPS z avtomatskim certifikatom (Let's Encrypt)",
    "Dodaj input validacijo z decoded JSON schema (npr. go-playground/validator)",
    "Implementiraj refresh token mehanizem za daljše seje",
    "Dodaj strukturirano logiranje z zahtevo ID (correlation ID)",
    "Pripravi threat model in varnostni pregled",
    "Razmisli o podpori za Linux/Mac pakete (ne samo Windows)",
])

add_heading("11.4 Dolgoročno", level=2)
add_bullets([
    "Razmisli o prehodu na mikrostoritve, če bo aplikacija rasla",
    "Implementiraj backup in restore strategijo za MongoDB",
    "Dodaj observability (OpenTelemetry, ki je že v odvisnostih)",
    "Implementiraj večjezično podporo z dodajanjem jezikov",
    "Razmisli o certificiranju za PCI-DSS, če bo procesiral plačilne kartice",
])

# ============================================================================
# 12. SKLEP
# ============================================================================
add_heading("12. Sklepna ocena", level=1)

add_para(
    "NutrixPOS je obetajoč odprtokodni POS sistem s sodobno arhitekturo "
    "in dobrim izborom tehnologij. Modularni pristop z vtičniškim sistemom "
    "za module je dobra osnova za nadaljnjo rast. Trenutno je projekt "
    "v zgodnji razvojni fazi (v0.5.5) in še ni primeren za produkcijsko "
    "uporabo brez dodatnega truda na področju varnosti in testiranja."
)

add_para("Ocene po kategorijah (1–5):", bold=True)

score_table = doc.add_table(rows=1, cols=3)
score_table.style = "Light Grid Accent 1"
hdr = score_table.rows[0].cells
hdr[0].text = "Kategorija"
hdr[1].text = "Ocena"
hdr[2].text = "Komentar"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(c, "1F3A5F")

scores = [
    ("Arhitektura", "4/5", "Modularna, razširljiva, dobro zasnovana"),
    ("Tehnološki sklad", "5/5", "Sodobne, pravilno izbrane tehnologije"),
    ("Varnost", "2/5", "Več kritičnih težav, zahteva takojšnje posredovanje"),
    ("Testiranje", "1/5", "Brez testov, vendar okvirji pripravljeni"),
    ("Kakovost kode", "3/5", "Srednje, z nekaj neskladji in tipkarskimi napakami"),
    ("Dokumentacija", "3/5", "AGENTS.md odličen, README minimalen"),
    ("CI/CD", "2/5", "Samo Windows paket, brez testov v CI"),
    ("Skupna ocena", "2.9/5", "Obetajoče, vendar še ni produkcijsko pripravljen"),
]
for s in scores:
    add_table_row(score_table, s)
    if s[0] == "Skupna ocena":
        for c in score_table.rows[-1].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph()

add_para(
    "Priporočilo: Pred produkcijsko uvedbo obvezno nasloviti vsaj "
    "kritične varnostne težave (JWT skrivnost, setup endpoint, CORS, "
    "MongoDB geslo) in dodati osnovne teste za kritične poti. Po teh "
    "popravkih lahko projekt ocenimo kot MVP-pripravljen za majhne "
    "scenarije (npr. ena lokacija, notranje omrežje). Za resnejše "
    "scenarije (več lokacij, javna izpostavljenost) so potrebne še "
    "dodatne izboljšave na področju varnosti in ops.",
    italic=True,
)

# --- Shrani ---
import os
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"OK: {OUT}")
